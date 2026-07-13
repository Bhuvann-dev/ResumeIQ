"""Render the improved resume (simple markdown) to a clean, ATS-friendly
DOCX or PDF.

Single column, standard fonts, real headings and bullet lists — exactly the
structure ATS parsers handle best. Deliberately forgiving about the markdown it
receives: smaller models emit inconsistent heading levels (`#`..`####`) and a
variety of bullet glyphs, so we normalize rather than drop lines. The PDF keeps
URLs and emails as clickable links.
"""

import io
import re

from docx import Document
from docx.shared import Pt
from fpdf import FPDF
from fpdf.enums import XPos, YPos

_BOLD = re.compile(r"\*\*(.+?)\*\*")
_ITALIC = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")
_CODE = re.compile(r"`(.+?)`")

# Any run of 1-6 leading '#'.
_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
# A URL or email token, for clickable links in the PDF export. The bare-domain
# case requires a "/path" so we linkify github.com/user but not Node.js.
_LINK = re.compile(
    r"^(https?://\S+|www\.\S+|[\w-]+(?:\.[\w-]+)+/\S*|[\w.+-]+@[\w-]+\.[\w.-]+)$"
)
# Common bullet glyphs models use, incl. the Unicode replacement char (mojibake).
_BULLET = re.compile(r"^[\-\*•·▪●‣⁃�]\s*(.*)$")


def _clean(text: str) -> str:
    """Drop inline markdown markers and stray mojibake — ATS parsers want plain text."""
    text = text.replace("�", "")
    text = _BOLD.sub(r"\1", text)
    text = _ITALIC.sub(r"\1", text)
    text = _CODE.sub(r"\1", text)
    return text.strip()


def markdown_to_docx(markdown: str) -> bytes:
    doc = Document()

    # Times New Roman across body and headings — clean, professional, ATS-safe.
    for style_name in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3", "List Bullet"):
        try:
            doc.styles[style_name].font.name = "Times New Roman"
        except KeyError:
            pass
    doc.styles["Normal"].font.size = Pt(11)

    for raw in markdown.splitlines():
        stripped = raw.strip()
        if not stripped:
            continue

        heading = _HEADING.match(stripped)
        if heading:
            hashes = len(heading.group(1))
            text = _clean(heading.group(2))
            # 1 '#' -> title (0), 2 -> section (1), 3+ -> sub (2)
            level = 0 if hashes == 1 else 1 if hashes == 2 else 2
            if text:
                doc.add_heading(text, level=level)
            continue

        bullet = _BULLET.match(stripped)
        if bullet:
            text = _clean(bullet.group(1))
            if text:
                doc.add_paragraph(text, style="List Bullet")
            continue

        text = _clean(stripped)
        if text:
            doc.add_paragraph(text)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# Common Unicode punctuation -> ASCII, so latin-1 core fonts don't render "?".
_TRANSLIT = {
    "–": "-", "—": "-",      # en / em dash
    "‘": "'", "’": "'",      # curly single quotes
    "“": '"', "”": '"',      # curly double quotes
    "…": "...",                    # ellipsis
    "•": "-", "·": "-",       # bullets that slip past _BULLET
    " ": " ",                      # non-breaking space
}


def _latin1(text: str) -> str:
    # Core PDF fonts are latin-1; transliterate common punctuation, then
    # replace anything still outside latin-1 rather than crash.
    for uni, ascii_ in _TRANSLIT.items():
        text = text.replace(uni, ascii_)
    return text.encode("latin-1", "replace").decode("latin-1")


def _render_contact(pdf: FPDF, text: str) -> None:
    """Center the contact line, rendering URLs/emails as clickable blue links."""
    pdf.set_font("Times", "", 9)
    total = pdf.get_string_width(_latin1(text))
    pdf.set_x(max(pdf.l_margin, (pdf.w - total) / 2))
    for token in re.split(r"(\s+)", text):
        bare = token.strip()
        if bare and _LINK.match(bare):
            if "@" in bare and not bare.startswith("http"):
                url = "mailto:" + bare
            elif bare.startswith(("http://", "https://")):
                url = bare
            else:  # www. or a bare domain/path
                url = "https://" + bare
            pdf.set_text_color(37, 99, 235)
            pdf.write(5, _latin1(bare), link=url)
            pdf.set_text_color(0, 0, 0)
        elif token:
            pdf.write(5, _latin1(token))
    pdf.ln(8)


def markdown_to_pdf(markdown: str) -> bytes:
    pdf = FPDF(format="A4")
    pdf.set_margins(18, 15, 18)
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    expect_contact = False
    for raw in markdown.splitlines():
        stripped = raw.strip()
        if not stripped:
            continue

        heading = _HEADING.match(stripped)
        if heading:
            text = _clean(heading.group(2))
            if not text:
                continue
            hashes = len(heading.group(1))
            if hashes == 1:  # name / title
                pdf.set_font("Times", "B", 20)
                pdf.cell(0, 10, _latin1(text), align="C", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                expect_contact = True
            elif hashes == 2:  # section
                pdf.ln(3)
                pdf.set_font("Times", "B", 12)
                pdf.cell(0, 6, _latin1(text.upper()), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                y = pdf.get_y()
                pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
                pdf.ln(1)
                expect_contact = False
            else:  # sub-heading (job/project title)
                pdf.ln(1)
                pdf.set_font("Times", "B", 10)
                pdf.multi_cell(0, 5, _latin1(text), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                expect_contact = False
            continue

        bullet = _BULLET.match(stripped)
        if bullet:
            text = _clean(bullet.group(1))
            if text:
                pdf.set_font("Times", "", 10)
                pdf.set_x(pdf.l_margin + 4)
                pdf.multi_cell(0, 5, _latin1("-  " + text), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            expect_contact = False
            continue

        text = _clean(stripped)
        if not text:
            continue
        if expect_contact:
            _render_contact(pdf, text)
            expect_contact = False
        else:
            pdf.set_font("Times", "", 10)
            pdf.multi_cell(0, 5, _latin1(text), new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    return bytes(pdf.output())
